import os
import logging
from flask import Flask, render_template, request, jsonify, session
from werkzeug.utils import secure_filename
from docx import Document
import uuid
import re
from dotenv import load_dotenv
load_dotenv()
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_pinecone import PineconeVectorStore
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import Document as LangchainDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.prompts import ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate
from pinecone import Pinecone, ServerlessSpec
import langsmith

from flask_cors import CORS

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app, resources={r"/*": {
    "origins": ["https://bents-model.vercel.app", "https://bents-model-4ppw.vercel.app"],
    "methods": ["GET", "POST", "OPTIONS", "PUT", "DELETE"],
    "allow_headers": ["Content-Type", "Authorization"],
    "supports_credentials": True
}})

logger.info("Flask app and CORS initialized")

app.secret_key = os.urandom(24)  # Set a secret key for sessions
logger.info("Session secret key set")

# Access your API keys (set these in environment variables)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
TRANSCRIPT_INDEX_NAMES = ["bents", "shop-improvement", "tool-recommendations"]
PRODUCT_INDEX_NAME = "bents-woodworking-products"
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_API_KEY"] = os.getenv("LANGSMITH_API_KEY")
os.environ["LANGCHAIN_ENDPOINT"] = "https://api.smith.langchain.com"
os.environ["LANGCHAIN_PROJECT"] = "jason-json"

logger.info("Environment variables and API keys loaded")

# Initialize Langchain components
embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
llm = ChatOpenAI(openai_api_key=OPENAI_API_KEY, model="gpt-4o-mini", temperature=0)
logger.info("Langchain components initialized")

# Initialize Pinecone
pc = Pinecone(api_key=PINECONE_API_KEY)
logger.info("Pinecone initialized")

# Create or connect to the Pinecone indexes
for INDEX_NAME in TRANSCRIPT_INDEX_NAMES + [PRODUCT_INDEX_NAME]:
    if INDEX_NAME not in pc.list_indexes().names():
        logger.info(f"Creating new Pinecone index: {INDEX_NAME}")
        pc.create_index(
            name=INDEX_NAME,
            dimension=1536,  # OpenAI embeddings dimension
            metric='cosine',
            spec=ServerlessSpec(cloud='aws', region='us-east-1')
        )
    else:
        logger.info(f"Connected to existing Pinecone index: {INDEX_NAME}")

# Create VectorStores
transcript_vector_stores = {name: PineconeVectorStore(index=pc.Index(name), embedding=embeddings, text_key="text") for name in TRANSCRIPT_INDEX_NAMES}
product_vector_store = PineconeVectorStore(index=pc.Index(PRODUCT_INDEX_NAME), embedding=embeddings, text_key="tags")
logger.info("VectorStores created")

# System instructions
SYSTEM_INSTRUCTIONS = """You are an AI assistant specialized in information retrieval from text documents.
        Always provide your responses in English, regardless of the language of the input or context.
        When given a document and a query:
        1. Analyze the document content and create an efficient index of key terms, concepts, and their locations within the text.
        2. When a query is received, use the index to quickly locate relevant sections of the document.
        3. Extract the most relevant information from those sections to form a concise and accurate answer.
        4. Always include the exact relevant content from the document, starting from the beginning of the relevant section. Use quotation marks to denote direct quotes.
        5. If applicable, provide a timestamp or location reference for where the information was found in the original document.
        6. After providing the direct quote, summarize or explain the answer if necessary.
        7. If the query cannot be answered from the given document, state this clearly.
        8. Always prioritize accuracy over speed. If you're not certain about an answer, say so.
        9. For multi-part queries, address each part separately and clearly.
        10. Aim to provide responses within seconds, even for large documents.
        11. Provide the timestamp for where the information was found in the original video. Use the format {{timestamp:MM:SS}} for timestamps under an hour, and {{timestamp:HH:MM:SS}} for longer videos.
        12. Do not include any URLs in your response. Just provide the timestamps in the specified format.
        13. When referencing timestamps that may be inaccurate, you can use language like "around", "approximately", or "in the vicinity of" to indicate that the exact moment may vary slightly.
        Remember, always respond in English, even if the query or context is in another language.
        """
logger.info("System instructions set")

def add_product(title, tags, link):
    logger.info(f"Adding product: {title}")
    product_id = str(uuid.uuid4())
    tags_text = ', '.join(tags) if isinstance(tags, list) else tags
    
    metadata = {
        "title": title,
        "tags": tags_text,
        "link": link
    }
    
    product_vector_store.add_texts([tags_text], metadatas=[metadata], ids=[product_id])
    logger.info(f"Product added with ID: {product_id}")
    return product_id

import base64
def get_all_products():
    logger.info("Starting get_all_products()")
    try:
        # Use the Pinecone index directly to fetch all vectors
        index = pc.Index(PRODUCT_INDEX_NAME)
        vector_dim = index.describe_index_stats()['dimension']
        zero_vector = [0.0] * vector_dim
        # Fetch all vectors (adjust limit if necessary)
        results = index.query(vector=zero_vector, top_k=10000, include_metadata=True)
        logger.info(f"Retrieved {len(results['matches'])} results from Pinecone")
        products = []
        for i, match in enumerate(results['matches']):
            logger.debug(f"Processing match {i+1}: ID: {match['id']}, Metadata: {match['metadata']}")
            metadata = match['metadata']
            image_data = metadata.get('image_data', '')
            image_url = ''
            if image_data:
                # Assuming the image is stored as base64 encoded JPEG
                image_url = f"data:image/jpeg;base64,{image_data}"
            product = [
                str(match['id']),  # Use Pinecone ID as the first element
                str(metadata.get('title', 'No Title')),
                str(metadata.get('tags', 'No Tags')),
                str(metadata.get('link', 'No Link')),
                image_url  # Add image URL as the fifth element
            ]
            products.append(product)
        logger.info(f"Final products list: {products}")
        return products
    except Exception as e:
        logger.error(f"Error in get_all_products: {str(e)}", exc_info=True)
        raise  # Re-raise the exception to be caught by the route handler

def process_answer(answer, url):
    logger.info("Processing answer")
    def replace_timestamp(match):
        timestamp = match.group(1)
        full_url = combine_url_and_timestamp(url, timestamp)
        return f"[video]({full_url})"
    
    processed_answer = re.sub(r'\{timestamp:([^\}]+)\}', replace_timestamp, answer)
    
    video_links = re.findall(r'\[video\]\(([^\)]+)\)', processed_answer)
    video_dict = {f'[video{i}]': link for i, link in enumerate(video_links)}
    
    for i, (placeholder, link) in enumerate(video_dict.items()):
        processed_answer = processed_answer.replace(f'[video]({link})', placeholder)
    
    logger.info("Answer processed")
    return processed_answer, video_dict

def combine_url_and_timestamp(base_url, timestamp):
    logger.info(f"Combining URL and timestamp: {base_url}, {timestamp}")
    parts = timestamp.split(':')
    if len(parts) == 2:
        minutes, seconds = map(int, parts)
        total_seconds = minutes * 60 + seconds
    elif len(parts) == 3:
        hours, minutes, seconds = map(int, parts)
        total_seconds = hours * 3600 + minutes * 60 + seconds
    else:
        logger.error(f"Invalid timestamp format: {timestamp}")
        raise ValueError("Invalid timestamp format")

    if '?' in base_url:
        return f"{base_url}&t={total_seconds}"
    else:
        return f"{base_url}?t={total_seconds}"    

def delete_product(product_id):
    logger.info(f"Deleting product: {product_id}")
    product_vector_store.delete([product_id])
    logger.info(f"Product deleted: {product_id}")

def update_product(product_id, title, tags, link):
    logger.info(f"Updating product: {product_id}")
    delete_product(product_id)
    new_id = add_product(title, tags, link)
    logger.info(f"Product updated: {product_id} -> {new_id}")

def extract_text_from_docx(file):
    logger.info(f"Extracting text from DOCX: {file.filename}")
    doc = Document(file)
    text = "\n".join([para.text for para in doc.paragraphs])
    logger.info("Text extracted from DOCX")
    return text

def extract_metadata_from_text(text):
    logger.info("Extracting metadata from text")
    title = text.split('\n')[0] if text else "Untitled Video"
    logger.info(f"Extracted title: {title}")
    return {"title": title}

def upsert_transcript(transcript_text, metadata, index_name):
    logger.info(f"Upserting transcript to index: {index_name}")
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    chunks = text_splitter.split_text(transcript_text)
    
    documents = []
    for i, chunk in enumerate(chunks):
        chunk_metadata = metadata.copy()
        chunk_metadata['chunk_id'] = f"{metadata['title']}_chunk_{i}"
        chunk_metadata['url'] = metadata.get('url', '')
        documents.append(LangchainDocument(page_content=chunk, metadata=chunk_metadata))
    
    transcript_vector_stores[index_name].add_documents(documents)
    logger.info(f"Transcript upserted to index: {index_name}")

@app.route('/')
@app.route('/database')
def serve_spa():
    logger.info("Serving SPA")
    return render_template('index.html')

@app.route('/chat', methods=['POST'])
def chat():
    logger.info("Chat endpoint called")
    try:
        data = request.json
        logger.debug(f"Received data: {data}")

        user_query = data['message']
        selected_index = data['selected_index']
        chat_history = data.get('chat_history', [])

        logger.debug(f"Chat history received: {chat_history}")

        # Format chat history for ConversationalRetrievalChain
        formatted_history = []
        for i in range(0, len(chat_history) - 1, 2):
            human = chat_history[i]
            ai = chat_history[i + 1] if i + 1 < len(chat_history) else ""
            formatted_history.append((human, ai))

        logger.debug(f"Formatted chat history: {formatted_history}")

        retriever = transcript_vector_stores[selected_index].as_retriever(search_kwargs={"k": 3})
        
        prompt = ChatPromptTemplate.from_messages([
            SystemMessagePromptTemplate.from_template(SYSTEM_INSTRUCTIONS),
            HumanMessagePromptTemplate.from_template("Context: {context}\n\nChat History: {chat_history}\n\nQuestion: {question}")
        ])
        
        qa_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            combine_docs_chain_kwargs={"prompt": prompt},
            return_source_documents=True
        )
        
        logger.info("Querying ConversationalRetrievalChain")
        result = qa_chain({"question": user_query, "chat_history": formatted_history})
        
        initial_answer = result['answer']
        context = [doc.page_content for doc in result['source_documents']]
        url = result['source_documents'][0].metadata.get('url', '') if result['source_documents'] else None
        
        logger.debug(f"Initial answer: {initial_answer}")
        logger.debug(f"URL: {url}")
        
        # Process the answer to replace timestamps and extract video links
        processed_answer, video_dict = process_answer(initial_answer, url)
        
        logger.debug(f"Processed answer: {processed_answer}")
        logger.debug(f"Video dict: {video_dict}")
        
        # Get embedding for the processed answer
        answer_embedding = embeddings.embed_query(processed_answer)
        
        logger.debug(f"Generated answer embedding. Shape: {len(answer_embedding)}")
        
        # Use answer embedding for product search
        product_index = pc.Index(PRODUCT_INDEX_NAME)
        logger.info(f"Querying product index: {PRODUCT_INDEX_NAME}")
        
        try:
            product_results = product_index.query(
                vector=answer_embedding,
                top_k=5,  # Retrieve more products to ensure we get all
                include_metadata=True
            )
            logger.debug(f"Product search results: {product_results}")
        except Exception as e:
            logger.error(f"Error querying product index: {str(e)}", exc_info=True)
            product_results = {'matches': []}
        
        matching_products = []
        non_matching_products = []
        
        for match in product_results['matches']:
            logger.debug(f"Processing match: {match}")
            product = {
                'title': match['metadata'].get('title', 'Untitled'),
                'tags': match['metadata'].get('tags', ''),
                'link': match['metadata'].get('link', '')
            }
            
            tags = match['metadata'].get('tags', '').split(',')
            logger.debug(f"Tags for product: {tags}")
            
            if any(tag.strip().lower() in processed_answer.lower() for tag in tags):
                matching_products.append(product)
                logger.debug(f"Added matching product: {product}")
